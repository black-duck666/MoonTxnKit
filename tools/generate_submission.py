from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission"
OUT.mkdir(parents=True, exist_ok=True)
PDF = OUT / "MoonTxnKit_项目申报书_复审版.pdf"
DOCX = OUT / "MoonTxnKit_项目申报书_复审版.docx"
MD = OUT / "MoonTxnKit_项目申报书_复审版.md"
CHECKLIST = OUT / "MoonTxnKit_复审检查清单.md"
EXTERNAL = ROOT.parent / "MoonTxnKit_提交材料_最新版"

TITLE = "MoonTxnKit：面向 MoonBit 的确定性状态事务与恢复内核"
GITHUB = "https://github.com/black-duck666/MoonTxnKit"
GITLINK = "https://www.gitlink.org.cn/black-duck/MoonTxnKit"
ACCENT_HEX = "285B75"
ACCENT = colors.HexColor("#285B75")
INK = colors.HexColor("#202A33")
MUTED = colors.HexColor("#63717C")
LIGHT = colors.HexColor("#EDF4F7")

BASIC = [
    ("项目名称", TITLE),
    ("参赛者", "马识途"),
    ("联系方式", "【提交前填写：手机、邮箱】"),
    ("GitHub 仓库", GITHUB),
    ("GitLink 仓库", GITLINK),
    ("项目方向", "MoonBit 基础生态库 / 原子状态迁移 / MVCC 与恢复"),
    ("是否为移植项目", "否"),
    ("当前基础", "21+ 次有效提交；23 项测试；四后端与 GitHub CI 通过"),
]

PAGE1 = [
    (
        "二、项目定位与真正价值",
        [
            "MoonTxnKit 不是 transactions 工作流程的教学玩具，也不试图再造 SQL 数据库。项目定位在“直接修改 Map”与“引入完整数据库”之间：为工作流调度器、规则引擎、模拟器、内存服务和测试替身提供可嵌入的确定性状态提交内核。",
            "这类组件经常需要先检查多个业务事实，再让一批状态全部生效或全部不生效，同时能够解释并发冲突、记录提交结果并在重启后恢复。如果各项目自行实现，会重复出现版本号、临时副本、补偿代码、锁和日志格式。MoonTxnKit 把这组共性收敛为稳定的 MoonBit API，且不绑定磁盘、线程、网络和平台 IO。",
        ],
    ),
    (
        "三、目标用户与复用场景",
        [
            "工作流与任务调度：任务仍为 queued 且 owner 不存在时，才能原子写入 running 与 owner，多个 worker 竞争时最多一个成功。",
            "库存和资源预占：库存、订单状态与幂等标记在同一快照检查，扣减、状态迁移和预占记录使用同一提交版本。",
            "规则与策略引擎：规则读取的一组事实在结论提交前不能被其他事务改变，Serializable 读集校验负责阻断过期结论。",
            "模拟器与测试替身：相同命令序列在不同后端得到相同版本、冲突和恢复结果，适合差异测试与回放。",
            "轻量持久化工具：核心生成可校验逻辑 WAL，上层可选择文件、浏览器存储、对象存储或数据库作为保存位置。",
        ],
    ),
    (
        "四、合理抽象：AtomicPlan",
        [
            "AtomicPlan 用 condition 描述提交前必须成立的业务事实，用 write 描述必须一起生效的状态集合。Engine::execute 在同一 Serializable 快照读取全部条件；条件不满足返回 ConditionFailed，并发验证失败返回 CommitRejected，两类失败具有不同语义和稳定 JSON 证据。",
            "上层只需要表达“什么必须仍然成立”和“哪些状态必须一起变化”，无需理解版本链、保存点或 WAL。任一条件失败时版本不会推进，也不会留下部分写入。该抽象足以表达任务领取、幂等消费、订单状态机、资源租约和策略配置切换，又没有把业务对象强制建模成数据库表。",
        ],
    ),
    (
        "五、项目边界与生态关系",
        [
            "项目不实现 SQL、索引、查询优化、文件刷盘、网络复制和分布式共识，也不宣称替代成熟数据库。它服务于“不值得引入完整数据库，但不能接受部分状态和不可恢复更新”的 MoonBit 组件。",
            "2026 年 6 月 30 日以 MVCC、snapshot isolation、serializable transaction、write-ahead log、transactional key-value 等关键词检查 Mooncakes。moonbitlang/async 的 retry 处理任务失败重试，不维护多版本状态；Lampese/lomo 的事务服务于 Loro CRDT 文档模型。未发现以通用原子状态计划、MVCC 隔离和逻辑恢复为组合定位的 MoonBit 包。",
        ],
    ),
]

PAGE2 = [
    (
        "六、方案位置对比",
        [
            "直接 Map 更新：依赖最少，但没有一致快照、整批原子性、并发冲突证据和恢复协议，复杂状态迁移最终会散落补偿与回滚代码。",
            "完整数据库：能力全面，但会引入 SQL、存储格式、运行时和部署约束；模拟器、规则引擎、浏览器工具和测试替身往往不需要这些重量。",
            "MoonTxnKit：只提供状态事务、冲突和逻辑恢复，保持纯 MoonBit 与后端中立；上层保留领域模型、编码、执行器和持久化选择。",
        ],
    ),
    (
        "七、核心技术设计",
        [
            "引擎为每个键维护按提交版本递增的版本链。事务固定 snapshot_version，读取不超过快照版本的最新值；待提交写在事务内优先可见。",
            "Snapshot Isolation 使用 first-committer-wins 检查写写冲突；Serializable 模式额外记录并验证读集，阻止检查依据在提交前被并发事务改变。整批写共享一个提交版本。",
            "每个写事务生成带校验值的逻辑 WAL。恢复先验证校验值、提交版本连续性和重复记录一致性，再应用数据；相同记录可幂等跳过，损坏或分叉记录会被拒绝。",
            "历史压缩以最老活跃快照为低水位，保留该快照可见锚点和其后版本，使长读事务在压缩前后仍看到同一状态。",
        ],
    ),
    (
        "八、创新点",
        [
            "把业务前置条件与底层事务控制解耦，形成可被多个领域直接调用的 AtomicPlan，而不是只暴露 begin/put/commit 教学流程。",
            "把隔离、冲突、日志、恢复和低水位压缩组织成无平台依赖的纯状态内核，允许不同后端自由选择持久化与执行方式。",
            "业务条件失败、并发冲突、WAL 损坏和恢复结果均使用结构化类型与稳定 JSON，便于 CI、调试器和可视化工具消费。",
            "核心不读取系统时钟、不使用随机数，同一命令序列可以在 Native、JavaScript、Wasm 和 Wasm-GC 后端进行差异验证。",
        ],
    ),
    (
        "九、当前成果与可运行证据",
        [
            "仓库已形成 21+ 次有效提交，覆盖 MVCC 引擎、两种隔离级别、保存点、AtomicPlan、WAL、恢复、压缩、JSON 报告、CLI、文档和持续集成。",
            "23 项确定性测试覆盖快照重复读、写写冲突、写偏斜阻断、保存点回滚、任务原子领取、库存预占、幂等消费、WAL 校验、幂等重放、损坏拒绝和低水位压缩；四个编译后端全部通过。",
            "CLI 同时演示银行转账与工作流任务领取：第一个 worker 原子写入 running 和 owner，第二个 worker 收到 condition_failed；随后五条 WAL 在新引擎中恢复，键值、历史版本和统计一致。",
        ],
    ),
    (
        "十、已知限制",
        [
            "当前键和值采用 String，尚未提供泛型编码适配；Serializable 只验证点读集合，尚无范围扫描和谓词冲突；逻辑 WAL 不直接承诺文件 fsync 或跨进程崩溃原子性。这些限制在 README 和设计文档中明确记录。",
        ],
    ),
    (
        "十一、后续计划与验收",
        [
            "0.3：增加版本化 WAL 二进制编解码、Bytes/领域编码适配和排序模型随机对照测试。",
            "0.4：探索范围索引、谓词冲突跟踪和 Serializable Snapshot Isolation 依赖图。",
            "0.5：提供文件、浏览器存储和对象存储适配包，建立检查点与崩溃恢复演练。",
            "阶段验收继续以四后端测试、公开 CI、结构化差异输出和带日期的性能记录为准，不使用无法复现的性能宣传。",
        ],
    ),
]


def register_font():
    for candidate in [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("CN", str(candidate)))
            return "CN"
    return "Helvetica"


FONT = register_font()


def escape(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont(FONT, 7.6)
    canvas.setFillColor(MUTED)
    canvas.drawString(16 * mm, 8 * mm, "MoonTxnKit 项目申报书 - 复审版")
    canvas.drawRightString(A4[0] - 16 * mm, 8 * mm, f"第 {doc.page} 页 / 共 2 页")
    canvas.restoreState()


def build_pdf():
    doc = BaseDocTemplate(
        str(PDF),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=12 * mm,
        bottomMargin=14 * mm,
    )
    doc.addPageTemplates(
        [
            PageTemplate(
                id="proposal",
                frames=[
                    Frame(
                        doc.leftMargin,
                        doc.bottomMargin,
                        doc.width,
                        doc.height,
                        id="content",
                    )
                ],
                onPage=footer,
            )
        ]
    )
    title = ParagraphStyle(
        "title",
        fontName=FONT,
        fontSize=14.2,
        leading=18,
        textColor=ACCENT,
        alignment=1,
        spaceAfter=4,
    )
    subtitle = ParagraphStyle(
        "subtitle",
        fontName=FONT,
        fontSize=8.4,
        leading=10.5,
        textColor=MUTED,
        alignment=1,
        spaceAfter=5,
    )
    heading = ParagraphStyle(
        "heading",
        fontName=FONT,
        fontSize=10.5,
        leading=13,
        textColor=ACCENT,
        spaceBefore=3,
        spaceAfter=2,
    )
    body = ParagraphStyle(
        "body",
        fontName=FONT,
        fontSize=8.7,
        leading=11.9,
        textColor=INK,
        firstLineIndent=13,
        spaceAfter=1.7,
    )
    numbered = ParagraphStyle(
        "numbered",
        fontName=FONT,
        fontSize=8.55,
        leading=11.5,
        textColor=INK,
        leftIndent=12,
        firstLineIndent=-10,
        spaceAfter=1.3,
    )
    small = ParagraphStyle(
        "small",
        fontName=FONT,
        fontSize=7.8,
        leading=10,
        textColor=INK,
    )
    story = [
        Paragraph("MoonTxnKit 项目申报书", title),
        Paragraph("确定性状态事务与恢复内核 - 复审版", subtitle),
        Paragraph("一、基本信息", heading),
    ]
    table = Table(
        [[Paragraph(escape(k), small), Paragraph(escape(v), small)] for k, v in BASIC],
        colWidths=[31 * mm, 131 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), LIGHT),
                ("TEXTCOLOR", (0, 0), (0, -1), ACCENT),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D9E1")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 2.3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2.3),
            ]
        )
    )
    story.extend([table, Spacer(1, 2)])
    numbered_sections = {
        "三、目标用户与复用场景",
        "六、方案位置对比",
        "八、创新点",
        "十一、后续计划与验收",
    }
    for page_index, sections in enumerate([PAGE1, PAGE2]):
        if page_index == 1:
            story.append(PageBreak())
        for section_title, paragraphs in sections:
            story.append(Paragraph(escape(section_title), heading))
            use_numbers = section_title in numbered_sections
            for index, text in enumerate(paragraphs):
                prefix = f"{index + 1}. " if use_numbers else ""
                story.append(
                    Paragraph(escape(prefix + text), numbered if use_numbers else body)
                )
    doc.build(story)
    pages = len(PdfReader(str(PDF)).pages)
    if pages != 2:
        raise RuntimeError(f"expected 2 PDF pages, got {pages}")


def set_run_font(run, size=9.0, bold=False, color="202A33"):
    run.font.name = "Microsoft YaHei"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        fonts.set(qn(key), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def add_sections(doc, sections):
    numbered_sections = {
        "三、目标用户与复用场景",
        "六、方案位置对比",
        "八、创新点",
        "十一、后续计划与验收",
    }
    for section_title, paragraphs in sections:
        heading = doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.keep_with_next = True
        set_run_font(heading.add_run(section_title), 10.7, True, ACCENT_HEX)
        use_numbers = section_title in numbered_sections
        for text in paragraphs:
            paragraph = doc.add_paragraph(style="List Number" if use_numbers else None)
            paragraph.paragraph_format.space_after = Pt(2.4)
            paragraph.paragraph_format.line_spacing = 1.12
            paragraph.paragraph_format.first_line_indent = None if use_numbers else Cm(0.7)
            set_run_font(paragraph.add_run(text), 8.9)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(8.9)
    normal.paragraph_format.space_after = Pt(2.4)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, before, after in [
        ("Heading 1", 10.7, 5, 2),
        ("Heading 2", 9.8, 4, 2),
        ("Heading 3", 9.2, 3, 1),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT_HEX)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    set_run_font(title.add_run("MoonTxnKit 项目申报书"), 16.2, True, ACCENT_HEX)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    set_run_font(
        subtitle.add_run("确定性状态事务与恢复内核 - 复审版"),
        9.2,
        False,
        "63717C",
    )
    heading = doc.add_paragraph(style="Heading 1")
    set_run_font(heading.add_run("一、基本信息"), 10.7, True, ACCENT_HEX)
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    for key, value in BASIC:
        cells = table.add_row().cells
        for cell, width in zip(cells, [1700, 7900]):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        shade(cells[0], "EDF4F7")
        set_run_font(cells[0].paragraphs[0].add_run(key), 8.3, True, ACCENT_HEX)
        set_run_font(cells[1].paragraphs[0].add_run(value), 8.3)

    add_sections(doc, PAGE1)
    doc.add_page_break()
    add_sections(doc, PAGE2)
    footer_paragraph = doc.sections[0].footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(
        footer_paragraph.add_run("MoonTxnKit 项目申报书 - 复审版  |  共 2 页"),
        7.5,
        False,
        "63717C",
    )
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "MoonBit 国产开源生态大赛复审材料"
    doc.core_properties.author = "MoonTxnKit contributors"
    doc.save(DOCX)


def build_text_files():
    lines = [f"# {TITLE}", "", "## 一、基本信息", ""]
    lines.extend([f"- {key}：{value}" for key, value in BASIC])
    for section_title, paragraphs in PAGE1 + PAGE2:
        lines.extend(["", f"## {section_title}", ""])
        lines.extend(paragraphs)
    MD.write_text("\n\n".join(lines), encoding="utf-8")
    CHECKLIST.write_text(
        "\n".join(
            [
                "# MoonTxnKit 复审检查清单",
                "",
                f"- [x] GitHub：{GITHUB}",
                f"- [x] GitLink：{GITLINK}",
                "- [x] GitLink 可公开克隆且与 GitHub 历史一致",
                "- [x] AtomicPlan 声明式原子状态迁移已实现",
                "- [x] 工作流、库存、幂等和删除场景测试齐备",
                "- [x] 23 项测试在四后端通过",
                "- [x] README、价值说明、架构、隔离、恢复和查重文档齐备",
                "- [x] 两页复审版 PDF 与 DOCX 已生成",
                "- [ ] 填写参赛者手机和邮箱",
                "- [ ] 在报名表上传复审版并重新提交",
            ]
        ),
        encoding="utf-8",
    )


def copy_external():
    EXTERNAL.mkdir(parents=True, exist_ok=True)
    for path in (PDF, DOCX, MD, CHECKLIST):
        shutil.copy2(path, EXTERNAL / path.name)


if __name__ == "__main__":
    build_pdf()
    build_docx()
    build_text_files()
    copy_external()
    print(PDF)
    print(DOCX)
